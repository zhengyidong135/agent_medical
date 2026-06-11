from __future__ import annotations

import io
import json
import re
import shutil
import uuid
import zipfile
import zlib
from datetime import datetime
from html import unescape
from pathlib import Path
from typing import Any

from config import (
    RAG_FILES_DIR,
    RAG_TEXT_DIR,
    SUPPORTED_RAG_EXTENSIONS,
    load_rag_index,
    save_rag_index,
    ensure_rag_store,
    safe_filename,
)

def read_text_file(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        if parts:
            return "\n".join(parts)
    except Exception:
        pass

    with zipfile.ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
        text_parts = []
        for name in names:
            xml = archive.read(name).decode("utf-8", errors="ignore")
            paragraphs = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, flags=re.DOTALL)
            text_parts.extend(unescape(part) for part in paragraphs)
        return "\n".join(text_parts)


def decode_pdf_literal(value: str) -> str:
    value = value.replace(r"\(", "(").replace(r"\)", ")").replace(r"\\", "\\")
    value = re.sub(r"\\[nrtbf]", " ", value)
    return value


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- 第 {index} 页 ---\n{text}")
        if pages:
            return "\n\n".join(pages)
    except Exception:
        pass

    data = path.read_bytes()
    chunks: list[bytes] = []
    for match in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", data, flags=re.DOTALL):
        stream = match.group(1).strip()
        try:
            chunks.append(zlib.decompress(stream))
        except zlib.error:
            chunks.append(stream)
    if not chunks:
        chunks = [data]

    parts: list[str] = []
    for chunk in chunks:
        text = chunk.decode("latin-1", errors="ignore")
        parts.extend(decode_pdf_literal(item) for item in re.findall(r"\((.*?)\)", text, flags=re.DOTALL))
        hex_strings = re.findall(r"<([0-9A-Fa-f]{4,})>", text)
        for item in hex_strings:
            try:
                raw = bytes.fromhex(item)
                parts.append(raw.decode("utf-16-be", errors="ignore") or raw.decode("latin-1", errors="ignore"))
            except ValueError:
                continue
    return "\n".join(part.strip() for part in parts if part.strip())


def extract_upload_text(path: Path, extension: str) -> str:
    if extension in {".txt", ".md"}:
        return read_text_file(path)
    if extension == ".docx":
        return extract_docx_text(path)
    if extension == ".pdf":
        return extract_pdf_text(path)
    if extension == ".doc":
        return read_text_file(path)
    raise ValueError("不支持的文件格式。")


def chunk_text(text: str, chunk_size: int = 900, overlap: int = 160) -> list[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    chunks = []
    start = 0
    while start < len(normalized):
        end = min(start + chunk_size, len(normalized))
        chunks.append(normalized[start:end])
        if end == len(normalized):
            break
        start = max(0, end - overlap)
    return chunks


def tokenize_query(text: str) -> set[str]:
    ascii_terms = re.findall(r"[A-Za-z0-9_]{2,}", text.lower())
    chinese_terms = re.findall(r"[\u4e00-\u9fff]{2,}", text)
    chinese_bigrams = []
    for term in chinese_terms:
        chinese_bigrams.extend(term[index : index + 2] for index in range(max(0, len(term) - 1)))
    return set(ascii_terms + chinese_terms + chinese_bigrams)


def retrieve_rag_context(file_ids: list[str], query: str, top_k: int = 5) -> tuple[str, list[dict[str, Any]]]:
    if not file_ids:
        return "", []
    index = {item["id"]: item for item in load_rag_index()}
    terms = tokenize_query(query)
    scored: list[tuple[int, dict[str, Any], int, str]] = []
    fallback: list[tuple[int, dict[str, Any], int, str]] = []

    for file_id in file_ids:
        item = index.get(file_id)
        if not item:
            continue
        text_path = RAG_TEXT_DIR / f"{file_id}.txt"
        if not text_path.exists():
            continue
        chunks = chunk_text(text_path.read_text(encoding="utf-8", errors="ignore"))
        for chunk_index, chunk in enumerate(chunks[:2]):
            fallback.append((0, item, chunk_index, chunk))
        for chunk_index, chunk in enumerate(chunks):
            lower = chunk.lower()
            score = sum(1 for term in terms if term.lower() in lower)
            if not terms and chunk.strip():
                score = 1
            if score > 0:
                scored.append((score, item, chunk_index, chunk))

    scored.sort(key=lambda row: row[0], reverse=True)
    selected = scored[:top_k] if scored else fallback[:top_k]
    sources = [
        {
            "file_id": item["id"],
            "filename": item["filename"],
            "chunk": chunk_index + 1,
            "score": score,
            "fallback": score == 0,
        }
        for score, item, chunk_index, _chunk in selected
    ]
    context_parts = [
        f"[来源: {item['filename']} / 片段 {chunk_index + 1}]\n{chunk}"
        for score, item, chunk_index, chunk in selected
    ]
    return "\n\n".join(context_parts), sources


def save_uploaded_rag_file(filename: str, source: Any) -> dict[str, Any]:
    ensure_rag_store()
    original_name = safe_filename(filename)
    extension = Path(original_name).suffix.lower()
    if extension not in SUPPORTED_RAG_EXTENSIONS:
        raise ValueError("只支持 txt、md、pdf、docx、doc 文件。")

    file_id = uuid.uuid4().hex
    stored_name = f"{file_id}{extension}"
    file_path = RAG_FILES_DIR / stored_name
    with file_path.open("wb") as target:
        shutil.copyfileobj(source, target)

    text = extract_upload_text(file_path, extension).strip()
    if not text:
        file_path.unlink(missing_ok=True)
        raise ValueError("没有从文件中抽取到可检索文本。")

    chunks = chunk_text(text)
    text_path = RAG_TEXT_DIR / f"{file_id}.txt"
    text_path.write_text(text, encoding="utf-8")

    item = {
        "id": file_id,
        "filename": original_name,
        "extension": extension,
        "stored_name": stored_name,
        "text_length": len(text),
        "chunk_count": len(chunks),
        "created_at": datetime.now().isoformat(timespec="seconds"),
    }
    items = load_rag_index()
    items.append(item)
    save_rag_index(items)
    return item

def rebuild_rag_file_text(item: dict[str, Any]) -> dict[str, Any]:
    file_id = item["id"]
    source_path = RAG_FILES_DIR / item["stored_name"]
    if not source_path.exists():
        raise ValueError(f"源文件不存在: {item.get('filename', file_id)}")

    text = extract_upload_text(source_path, item["extension"]).strip()
    if not text:
        raise ValueError(f"没有从文件中抽取到可检索文本: {item.get('filename', file_id)}")

    chunks = chunk_text(text)
    (RAG_TEXT_DIR / f"{file_id}.txt").write_text(text, encoding="utf-8")
    item["text_length"] = len(text)
    item["chunk_count"] = len(chunks)
    item["rebuilt_at"] = datetime.now().isoformat(timespec="seconds")
    return item


def rebuild_all_rag_texts() -> list[dict[str, Any]]:
    items = load_rag_index()
    rebuilt = []
    for item in items:
        rebuilt.append(rebuild_rag_file_text(item))
    save_rag_index(items)
    return rebuilt


def get_rag_file(file_id: str) -> dict[str, Any]:
    item = next((entry for entry in load_rag_index() if entry.get("id") == file_id), None)
    if not item:
        raise ValueError("文件不存在。")
    return item


def get_rag_text(file_id: str) -> str:
    get_rag_file(file_id)
    text_path = RAG_TEXT_DIR / f"{file_id}.txt"
    if not text_path.exists():
        raise ValueError("解析文本不存在，请先重新解析文件。")
    return text_path.read_text(encoding="utf-8")


def build_rag_docx(file_id: str) -> bytes:
    item = get_rag_file(file_id)
    text = get_rag_text(file_id)
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("缺少 python-docx，无法导出 Word 文档。") from exc

    document = Document()
    document.add_heading(item.get("filename", "RAG document"), level=1)
    for block in text.splitlines():
        document.add_paragraph(block)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def delete_rag_file(file_id: str) -> dict[str, Any]:
    items = load_rag_index()
    item = next((entry for entry in items if entry.get("id") == file_id), None)
    if not item:
      raise ValueError("文件不存在。")

    stored_name = item.get("stored_name", "")
    if stored_name:
        (RAG_FILES_DIR / stored_name).unlink(missing_ok=True)
    (RAG_TEXT_DIR / f"{file_id}.txt").unlink(missing_ok=True)
    save_rag_index([entry for entry in items if entry.get("id") != file_id])
    return item
